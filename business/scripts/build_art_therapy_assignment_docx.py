from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "drafts" / "미디어소프트웨어학과_20211001_윤준하_미술심리치료_중간과제.md"
OUTPUT = ROOT / "drafts" / "미디어소프트웨어학과_20211001_윤준하_미술심리치료_중간과제.docx"
FONT_NAME = "Malgun Gothic"


def set_font(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.5)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0.7)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_font(run, 10.5)


def build() -> None:
    lines = INPUT.read_text(encoding="utf-8").splitlines()
    student_info = lines[0].strip()
    title = lines[2].lstrip("# ").strip()
    body_lines = [line.strip() for line in lines[4:] if line.strip()]

    doc = Document()
    configure_document(doc)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info.paragraph_format.space_after = Pt(16)
    run = info.add_run(student_info)
    set_font(run, 10.5)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(18)
    title_run = title_paragraph.add_run(title)
    set_font(title_run, 16, bold=True)

    buffer: list[str] = []
    for line in body_lines:
        if len(line) <= 34 and line.endswith(("다.", "다")) and buffer:
            buffer.append(line)
            add_body_paragraph(doc, " ".join(buffer))
            buffer = []
            continue

        if buffer:
            add_body_paragraph(doc, " ".join(buffer))
            buffer = []
        buffer.append(line)

    if buffer:
        add_body_paragraph(doc, " ".join(buffer))

    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build()

