from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
GENERATED = ROOT / "generated"
IMAGES = ROOT.parent / "images"

PLAN_SOURCE = GENERATED / "사업계획서_PULSE_2026_실제입력본.docx"
PLAN_OUTPUT = GENERATED / "사업계획서_PULSE_2026_실제입력본_사진삽입본.docx"
APPENDIX_SOURCE = GENERATED / "사업계획서_PULSE_2026_별첨2_증빙서류_입력본.docx"
APPENDIX_OUTPUT = GENERATED / "사업계획서_PULSE_2026_별첨2_증빙서류_입력본_사진삽입본.docx"


IMAGE_MAP = {
    "landing": IMAGES / "메인페이지.png",
    "service": IMAGES / "서비스대표사진.png",
    "persona": IMAGES / "MVP상세사진1-고객페르소나및여정지도.png",
    "reels": IMAGES / "MVP상세사진2-릴스생성.png",
    "influencer": IMAGES / "MVP상세사진3-인플루언서매칭.png",
    "architecture": IMAGES / "기술아키텍처.png",
    "business_model": IMAGES / "비즈니스모델.png",
}


def remove_all_runs(paragraph: Paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def set_caption_style(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in paragraph.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(102, 102, 102)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def replace_paragraph_with_picture(
    paragraph: Paragraph,
    image_path: Path,
    width: float,
    caption: str | None = None,
) -> None:
    remove_all_runs(paragraph)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    if caption:
        caption_para = insert_paragraph_after(paragraph, caption)
        set_caption_style(caption_para)


def replace_cell_with_picture(cell, image_path: Path, width: float, caption: str | None = None) -> None:
    cell.text = ""
    picture_paragraph = cell.paragraphs[0]
    picture_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    picture_paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    if caption:
        caption_paragraph = cell.add_paragraph(caption)
        set_caption_style(caption_paragraph)


def replace_cell_caption(cell, caption: str) -> None:
    cell.text = caption
    for paragraph in cell.paragraphs:
        set_caption_style(paragraph)


def add_section_title(doc: Document, title: str) -> Paragraph:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return paragraph


def add_section_note(doc: Document, text: str) -> Paragraph:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in paragraph.runs:
        run.font.size = Pt(10)
    return paragraph


def add_full_width_image(doc: Document, image_path: Path, caption: str, width: float = 6.0) -> None:
    picture_paragraph = doc.add_paragraph()
    picture_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    picture_paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    caption_paragraph = doc.add_paragraph(caption)
    set_caption_style(caption_paragraph)


def add_two_column_gallery(doc: Document, items: list[tuple[Path, str]], width: float = 2.75) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    for idx in range(0, len(items), 2):
        pair = items[idx : idx + 2]
        row = table.add_row()
        for col_idx, cell in enumerate(row.cells):
            if col_idx >= len(pair):
                cell.text = ""
                continue
            image_path, caption = pair[col_idx]
            replace_cell_with_picture(cell, image_path, width=width, caption=caption)


def build_plan_doc() -> Path:
    copy2(PLAN_SOURCE, PLAN_OUTPUT)
    doc = Document(str(PLAN_OUTPUT))

    # Summary table image row
    summary = doc.tables[4]
    replace_cell_with_picture(summary.cell(6, 1), IMAGE_MAP["landing"], width=1.35)
    replace_cell_with_picture(summary.cell(6, 2), IMAGE_MAP["service"], width=1.35)
    replace_cell_with_picture(summary.cell(6, 3), IMAGE_MAP["architecture"], width=1.30)
    replace_cell_with_picture(summary.cell(6, 4), IMAGE_MAP["business_model"], width=1.10)

    replace_cell_caption(summary.cell(7, 1), "서비스 랜딩 화면")
    replace_cell_caption(summary.cell(7, 2), "서비스 대표 대시보드")
    replace_cell_caption(summary.cell(7, 3), "기술 아키텍처")
    replace_cell_caption(summary.cell(7, 4), "비즈니스 모델")

    paragraph_map = {
        29: (IMAGE_MAP["service"], 6.1, "PULSE 대표 대시보드 화면"),
        57: (IMAGE_MAP["business_model"], 3.1, "PULSE 고객-이용-수익화 구조"),
        86: (IMAGE_MAP["influencer"], 6.1, "제휴 및 확장 기능 화면 예시"),
    }

    for paragraph_index, (image_path, width, caption) in paragraph_map.items():
        replace_paragraph_with_picture(doc.paragraphs[paragraph_index - 1], image_path, width, caption)

    doc.save(str(PLAN_OUTPUT))
    return PLAN_OUTPUT


def build_appendix_doc() -> Path:
    copy2(APPENDIX_SOURCE, APPENDIX_OUTPUT)
    doc = Document(str(APPENDIX_OUTPUT))

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_section_title(doc, "[붙임 7-2] MVP 상세 화면 캡처")
    add_section_note(doc, "현재 보유한 서비스 화면을 양식 흐름에 맞춰 보기 쉽게 정리하였습니다.")
    add_two_column_gallery(
        doc,
        [
            (IMAGE_MAP["service"], "서비스 대표 대시보드"),
            (IMAGE_MAP["persona"], "고객 페르소나 및 여정지도"),
            (IMAGE_MAP["reels"], "AI 홍보 영상 생성 화면"),
            (IMAGE_MAP["influencer"], "인플루언서 매칭 화면"),
        ],
    )

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_section_title(doc, "[붙임 7-3] 프로젝트 전체 기술 아키텍처")
    add_section_note(doc, "프론트엔드, 메인 백엔드, AI 서버, 데이터 저장 흐름을 한 장으로 확인할 수 있는 구조도입니다.")
    add_full_width_image(doc, IMAGE_MAP["architecture"], "PULSE 통합 기술 아키텍처", width=6.0)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_section_title(doc, "[붙임 7-4] 사용자 흐름 및 사업 모델 시각자료")
    add_section_note(doc, "서비스 진입 화면과 핵심 고객-사용-수익화 흐름을 함께 제시해 이해도를 높였습니다.")
    add_full_width_image(doc, IMAGE_MAP["landing"], "서비스 랜딩 및 진입 화면", width=6.0)
    add_full_width_image(doc, IMAGE_MAP["business_model"], "핵심 고객-서비스 이용-수익화 흐름도", width=3.0)

    doc.save(str(APPENDIX_OUTPUT))
    return APPENDIX_OUTPUT


def main() -> None:
    GENERATED.mkdir(parents=True, exist_ok=True)
    outputs = [build_plan_doc(), build_appendix_doc()]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
