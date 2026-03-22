from __future__ import annotations

from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent.parent
GENERATED = ROOT / "generated"
IMAGES = ROOT.parent / "images"


def find_generated(name_suffix: str) -> Path:
    return next(p for p in GENERATED.glob(f"*{name_suffix}") if "도식반영본" not in p.name)


PLAN_SOURCE = find_generated("실제입력본.docx")
APPENDIX_SOURCE = find_generated("별첨2_증빙서류_입력본.docx")
PLAN_OUTPUT = GENERATED / "사업계획서_PULSE_2026_실제입력본_도식반영본.docx"
APPENDIX_OUTPUT = GENERATED / "사업계획서_PULSE_2026_별첨2_증빙서류_입력본_도식반영본.docx"

IMAGE_MAP = {
    "kci": IMAGES / "KCI논문사진.png",
    "problem": IMAGES / "문제검증자료.png",
    "business_model": IMAGES / "비즈니스모델.png",
    "architecture": IMAGES / "기술아키텍처.png",
    "technical_diff": next(
        p for p in IMAGES.glob("*.png") if "차별성" in p.name
    ),
}


def remove_all_runs(paragraph: Paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def set_caption_style(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in paragraph.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(102, 102, 102)


def replace_paragraph_with_picture(paragraph: Paragraph, image_path: Path, width: float, caption: str) -> None:
    remove_all_runs(paragraph)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    caption_paragraph = insert_paragraph_after(paragraph, caption)
    set_caption_style(caption_paragraph)


def replace_cell_text(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(102, 102, 102)


def add_section_title(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(14)


def add_section_note(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in paragraph.runs:
        run.font.size = Pt(10)


def add_full_width_image(doc: Document, image_path: Path, caption: str | None = None, width: float = 6.0) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    if caption:
        caption_paragraph = doc.add_paragraph(caption)
        set_caption_style(caption_paragraph)


def build_plan_doc() -> Path:
    copy2(PLAN_SOURCE, PLAN_OUTPUT)
    doc = Document(str(PLAN_OUTPUT))

    # Summary table image row: keep only short references so body images stay within the 3-image rule.
    summary = doc.tables[4]
    replace_cell_text(summary.cell(6, 1), "문제 검증 자료\n(본문 4번 페이지 반영)")
    replace_cell_text(summary.cell(6, 2), "기술 차별성 자료\n(본문 5번 페이지 반영)")
    replace_cell_text(summary.cell(6, 3), "비즈니스모델 구체화\n(본문 8번 페이지 반영)")
    replace_cell_text(summary.cell(6, 4), "KCI 논문 등재 증빙\n(별첨 2 참조)")
    replace_cell_text(summary.cell(7, 1), "핵심 문제 구조")
    replace_cell_text(summary.cell(7, 2), "실행형 AI 루프")
    replace_cell_text(summary.cell(7, 3), "Free-Basic-Pro-B2B")
    replace_cell_text(summary.cell(7, 4), "평가 참고 증빙")

    paragraph_map = {
        29: (IMAGE_MAP["technical_diff"], 5.9, "PULSE 기술 차별성 설명 자료"),
        57: (IMAGE_MAP["business_model"], 5.4, "PULSE 비즈니스모델 구체화"),
        86: (IMAGE_MAP["problem"], 5.8, "외식업 자영업자 문제 검증 자료"),
    }

    for paragraph_index, (image_path, width, caption) in paragraph_map.items():
        replace_paragraph_with_picture(doc.paragraphs[paragraph_index - 1], image_path, width, caption)

    doc.save(str(PLAN_OUTPUT))
    return PLAN_OUTPUT


def build_appendix_doc() -> Path:
    copy2(APPENDIX_SOURCE, APPENDIX_OUTPUT)
    doc = Document(str(APPENDIX_OUTPUT))

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_section_title(doc, "[붙임 7-1] KCI 논문 등재 증빙")
    add_section_note(doc, "연구 기반 문제 정의와 분석 구조의 신뢰도를 보강하는 참고 증빙입니다.")
    add_full_width_image(doc, IMAGE_MAP["kci"], width=5.2)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_section_title(doc, "[붙임 7-2] 문제 검증 자료")
    add_section_note(doc, "외식업 자영업자가 실제로 겪는 실행 장벽을 한 장에서 이해할 수 있도록 정리한 도식입니다.")
    add_full_width_image(doc, IMAGE_MAP["problem"], width=6.0)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_section_title(doc, "[붙임 7-3] 기술 차별성 설명 자료")
    add_section_note(doc, "PULSE가 수집-이해-생성-실행을 하나의 루프로 연결하는 구조를 설명하는 참고 도식입니다.")
    add_full_width_image(doc, IMAGE_MAP["technical_diff"], width=4.8)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_section_title(doc, "[붙임 7-4] 프로젝트 전체 기술 아키텍처")
    add_section_note(doc, "프론트엔드, 메인 백엔드, AI 서버, 데이터 저장 흐름을 함께 보여주는 구조도입니다.")
    add_full_width_image(doc, IMAGE_MAP["architecture"], width=5.9)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_section_title(doc, "[붙임 7-5] 비즈니스모델 구체화")
    add_section_note(doc, "무료 진입, 유료 전환, 업셀, 제휴 확장까지 이어지는 사업 구조를 요약한 도식입니다.")
    add_full_width_image(doc, IMAGE_MAP["business_model"], width=5.4)

    doc.save(str(APPENDIX_OUTPUT))
    return APPENDIX_OUTPUT


def main() -> None:
    GENERATED.mkdir(parents=True, exist_ok=True)
    for output in [build_plan_doc(), build_appendix_doc()]:
        print(output)


if __name__ == "__main__":
    main()
