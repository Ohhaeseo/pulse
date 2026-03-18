from pathlib import Path
from shutil import copy2

from docx import Document


ROOT = Path(__file__).resolve().parent.parent
OFFICIAL = ROOT / "official_docs"
GENERATED = ROOT / "generated"
TEMPLATE = next(
    p for p in OFFICIAL.iterdir() if p.suffix.lower() == ".docx" and "[별첨 1]" in p.name
)
OUTPUT = GENERATED / "사업계획서_PULSE_2026_실제입력본.docx"


def set_paragraph(doc: Document, idx: int, text: str, style: str | None = None) -> None:
    p = doc.paragraphs[idx - 1]
    p.text = text
    if style:
        p.style = style


def fill_row(row, values):
    for i, value in enumerate(values):
        row.cells[i].text = value


def build() -> Path:
    GENERATED.mkdir(parents=True, exist_ok=True)
    if OUTPUT.exists():
        OUTPUT.unlink()
    copy2(TEMPLATE, OUTPUT)
    doc = Document(str(OUTPUT))
    tables = doc.tables

    # Table 4: General information
    general = tables[3]
    general.cell(0, 3).text = "네이버·카카오 리뷰 분석 기반 외식업 AI 마케팅 운영 플랫폼 PULSE"
    general.cell(1, 3).text = (
        "웹 서비스(1식), 네이버·카카오 리뷰 수집 및 분석 기능(1식), 고객 페르소나·고객 여정 도출 기능(1식), "
        "AI 숏폼 생성 기능(1식), 실행 제안형 대시보드 및 상권 분석 기능(1식)"
    )
    general.cell(2, 3).text = "대학생(4학년 재학) / 미디어소프트웨어학과 / AI·CTO"

    team_rows = [
        ["1", "공동대표", "마케팅·기획·디자인", "노태경 / 대학교 4학년 재학 / 서비스 기획, 마케팅 설계, 디자인 역량", "완료"],
        ["2", "공동대표", "프론트엔드 개발", "김혜린 / 대학교 4학년 재학 / React 기반 UI 구현 및 사용자 경험 설계 역량", "완료"],
        ["3", "공동대표", "백엔드 개발", "오해서 / 대학교 4학년 재학 / Spring Boot 기반 API 및 서버 연동 역량", "완료"],
    ]
    for row_idx, values in zip(range(5, 8), team_rows):
        row = general.rows[row_idx]
        row.cells[0].text = values[0]
        row.cells[1].text = values[1]
        row.cells[2].text = values[2]
        row.cells[3].text = values[2]
        row.cells[4].text = values[3]
        row.cells[5].text = values[3]
        row.cells[6].text = values[3]
        row.cells[7].text = values[4]

    # Table 5: Summary page
    summary = tables[4]
    summary.cell(0, 1).text = "PULSE"
    summary.cell(0, 3).text = "AI SaaS"
    summary.cell(1, 1).text = (
        "PULSE는 외식업 자영업자를 위한 AI 마케팅 운영 플랫폼으로, 회원가입 이후 리뷰 분석이 자동 시작되고 "
        "고객 이해부터 홍보 콘텐츠 제작, 다음 행동 제안까지 이어지는 실행형 서비스를 제공합니다."
    )
    summary.cell(2, 1).text = (
        "리뷰와 고객 데이터를 활용한 마케팅 의사결정은 중요하지만, 자영업자는 분석 역량, "
        "콘텐츠 제작 시간, 외부 대행 비용과 신뢰 문제로 실제 실행이 어렵습니다."
    )
    summary.cell(3, 1).text = (
        "프론트엔드, 메인 백엔드, AI 서버로 구성된 MVP 개발이 완료되었고, 네이버·카카오 리뷰 수집, "
        "토픽 분석, 고객 페르소나·고객 여정 생성, 대시보드와 영상 제작 UX까지 시연 가능한 상태입니다."
    )
    summary.cell(4, 1).text = (
        "파일럿 점포를 통한 PoC 운영으로 현장성을 검증한 뒤 구독형 SaaS 모델을 검증하고, "
        "향후 리뷰 운영, 인플루언서 매칭, 다점포 관리, B2B 제휴 영역으로 확장합니다."
    )
    summary.cell(5, 1).text = (
        "윤준하 대표자는 AI·CTO를 담당하고, 노태경(마케팅·기획·디자인), 김혜린(프론트엔드), "
        "오해서(백엔드)가 공동대표로 참여하는 역할 분담형 팀 구조로 MVP를 직접 구축했습니다."
    )
    summary.cell(7, 1).text = "[MVP개발1번째사진첨부: 메인 대시보드 또는 메인 화면]"
    summary.cell(7, 2).text = "[별첨2제출: 프로젝트 전체 기술적 아키텍처]"
    summary.cell(7, 3).text = "[별첨2제출: 사용자 입장 서비스 흐름도]"
    summary.cell(7, 4).text = "[별첨2제출: KCI 논문 등재 사진]"

    # Problem / Solution / Scale-up / Team narratives
    set_paragraph(doc, 18, "ㅇ 아이템 개요")
    set_paragraph(
        doc,
        19,
        "PULSE는 외식업 자영업자가 직접 마케팅을 공부하거나 별도 대행사를 쓰지 않아도, "
        "회원가입 이후 리뷰 분석이 자동 시작되고 고객 이해, 홍보 제작, 다음 행동 제안까지 이어지는 마케팅 실행 루프를 자동화하는 서비스입니다.",
    )
    set_paragraph(
        doc,
        20,
        "핵심은 단순 정보 제공이 아니라, 사장님이 바로 실행할 수 있도록 고객 인사이트와 콘텐츠 제작, 실행 제안까지 연결하는 데 있습니다.",
        "List Paragraph",
    )
    set_paragraph(doc, 22, "ㅇ 문제 인식 (Problem)")
    set_paragraph(
        doc,
        23,
        "외식업 자영업자는 현장 업무 비중이 높아 마케팅의 필요성을 인지하고도 실행까지 이어가기 어렵습니다.",
    )
    set_paragraph(
        doc,
        24,
        "특히 고객 선호 해석, 콘텐츠 기획, 홍보 이후 실행 의사결정에서 반복적인 어려움을 겪습니다.",
        "List Paragraph",
    )
    set_paragraph(
        doc,
        25,
        "기존 광고 플랫폼은 노출 중심이고, 일반 디자인·영상 툴은 제작 중심입니다. "
        "즉, 시장에는 분석과 제작, 실행을 하나의 흐름으로 묶는 통합형 솔루션이 부족합니다.",
    )
    set_paragraph(doc, 26, "ㅇ 실현 가능성 (Solution)")
    set_paragraph(
        doc,
        27,
        "PULSE는 React 프론트엔드, Spring Boot 메인 백엔드, FastAPI AI 서버로 구현되었으며, 회원가입 직후 리뷰 분석이 자동 시작되는 구조를 갖춘 MVP로 실제 사용자 흐름에 따라 시연이 가능합니다.",
    )
    set_paragraph(
        doc,
        28,
        "현재 네이버·카카오 리뷰 수집, 한국어 리뷰 토픽 분석, 고객 페르소나·고객 여정 도출, 행동 제안형 대시보드, 상권 분석, 홍보영상 생성 UX가 구현되어 있으며, 협약기간에는 영상 생성·리뷰 운영 API 연동과 PoC 운영을 통해 제품 완성도를 높일 계획입니다.",
        "List Paragraph",
    )
    set_paragraph(doc, 29, "[본문삽입: MVP 대표 화면 1장 - 메인 대시보드 또는 랜딩 화면]")
    set_paragraph(doc, 30, "[별첨2제출: MVP 상세 화면 - 리뷰 분석 및 페르소나 도출 화면]")
    set_paragraph(doc, 31, "[별첨2제출: MVP 상세 화면 - 홍보영상 생성 화면]")
    set_paragraph(doc, 32, "[별첨2제출: 프로젝트 전체 기술적 아키텍처 도표]")
    set_paragraph(doc, 33, "[별첨2제출: 사용자 입장 서비스 흐름도]")
    set_paragraph(doc, 34, "[별첨2제출: KCI 논문 등재 사진 - KCI 논문 검색 결과 또는 논문 첫 페이지 캡처]")
    set_paragraph(
        doc,
        35,
        "※ KCI 등재 논문: 「온라인 리뷰 데이터 기반 소상공인 전략 수립을 위한 고객 분석 프로세스 프레임워크 제안」, "
        "문화기술의 융합, 2026, 12권 1호, 367-380. 본 논문은 리뷰 기반 문제 정의와 고객 분석 프레임이 학술적으로도 구조화되었음을 보여줍니다.",
    )
    set_paragraph(
        doc,
        54,
        "ㅇ 사업화 추진 전략",
    )
    set_paragraph(
        doc,
        55,
        "초기 목표 고객은 SNS 마케팅 수요는 높지만 내부 인력과 분석 역량이 부족한 외식업 자영업자이며, 지역 단위 PoC로 사용성과 반복 사용 가능성을 먼저 검증합니다.",
    )
    set_paragraph(
        doc,
        56,
        "이후 구독형 SaaS로 전환하고, “리뷰 분석 + 숏폼 제작 + 실행 제안” 통합 가치로 초기 고객을 확보합니다. 기본 요금제는 분석·영상·대시보드 중심으로 구성하고, 향후에는 리뷰 운영, 인플루언서 매칭, 다점포 관리, B2B 제휴로 업셀링과 매출 단가 확대를 추진합니다.",
        "List Paragraph",
    )
    set_paragraph(doc, 57, "[본문삽입: 비즈니스 모델 도식 또는 사업 확장 로드맵 중 1개 선택]")
    set_paragraph(doc, 58, "[별첨2제출 가능: 사업 확장 로드맵 도표 또는 추가 설명 이미지]", "List Paragraph")
    set_paragraph(doc, 78, "ㅇ 경쟁사 분석 및 목표시장 진입 전략")
    set_paragraph(
        doc,
        79,
        "기존 디자인·영상 툴은 제작 중심이고, 광고 플랫폼은 집행 중심입니다. PULSE는 회원가입 직후 분석 시작, 리뷰 기반 고객 이해, 실행 제안형 대시보드까지 이어지는 운영 루프를 제공해 경쟁 서비스 대비 실행 연결성이 높습니다.",
    )
    set_paragraph(
        doc,
        80,
        "초기에는 MVP 시연과 파일럿 운영으로 레퍼런스를 확보하고, 이후 구독형 전환과 Pro 기능 확장으로 시장에 진입합니다.",
        "List Paragraph",
    )
    set_paragraph(doc, 84, "ㅇ 중장기 사업 확장 및 사회적 가치")
    set_paragraph(
        doc,
        85,
        "PULSE는 외식업 소상공인의 디지털 마케팅 정보 격차를 줄이고, 마케팅 실행 장벽을 낮추며, "
        "지역 상권 자생력 강화에 기여할 수 있습니다.",
    )
    set_paragraph(doc, 86, "[이미지삽입: 사업화 확장 단계 또는 사회적 가치 구조도]", "List Paragraph")
    set_paragraph(doc, 107, "ㅇ 대표자 및 팀원 보유 역량")
    set_paragraph(
        doc,
        108,
        "윤준하 대표자는 외식업 자영업자의 문제를 서비스 구조로 번역하고, 사용자가 실제로 행동할 수 있는 제품 흐름을 설계합니다.",
    )
    set_paragraph(
        doc,
        109,
        "특히 PULSE의 핵심 가치인 “분석 결과를 행동으로 연결하는 제품 구조”를 기획 단계부터 주도했으며, "
        "KCI 등재 논문을 통해 문제-해결 구조를 체계적으로 정리한 경험이 있습니다.",
        "List Paragraph",
    )
    set_paragraph(doc, 110, "[별첨2제출: 논문등재사진첨부 - KCI 논문 검색 결과 또는 논문 표지]")
    set_paragraph(doc, 111, "[별첨2제출 가능: 팀 역할 구조도]", "List Paragraph")

    # Schedule table
    schedule = tables[5]
    schedule_rows = [
        ["1", "핵심 분석 기능 고도화", "협약 1~2개월차", "리뷰 분석 정교화, 상권 분석 보강, UI 안정화"],
        ["2", "콘텐츠·운영 기능 고도화", "협약 2~4개월차", "영상 생성 품질 향상, 리뷰 운영 기능 및 API 연동 보강"],
        ["3", "PoC 점포 모집 및 운영", "협약 3~6개월차", "실제 외식업 점포 테스트 운영, 사용성 피드백 수집"],
        ["4", "서비스 검증 및 2단계 준비", "협약 5~8개월차", "성과 분석, 반복 사용성 검증, 사업화 완성도 보강"],
    ]
    for row_idx, values in zip(range(1, 5), schedule_rows):
        fill_row(schedule.rows[row_idx], values)

    # Budget tables
    budget1 = tables[6]
    budget1_rows = [
        ["외주용역비", "UI/브랜딩 고도화 및 디자인 자산 정리", "4,000,000"],
        ["지급수수료", "클라우드, API, 테스트 툴, 도메인 운영", "4,000,000"],
        ["재료비", "PoC 운영 및 시연·홍보 자료 제작", "2,000,000"],
        ["인건비", "협약 후 참여 인력 인건비 일부", "8,000,000"],
        ["지급수수료", "사용자 인터뷰 및 파일럿 운영 비용", "2,000,000"],
    ]
    for row_idx, values in zip(range(1, 6), budget1_rows):
        fill_row(budget1.rows[row_idx], values)
    fill_row(budget1.rows[6], ["합 계", "합 계", "20,000,000"])

    budget2 = tables[7]
    budget2_rows = [
        ["인건비", "서비스 운영 및 사업화 인력 보강", "8,000,000"],
        ["외주용역비", "제품 완성도 향상 및 세일즈 자산 보강", "4,000,000"],
        ["지급수수료", "SaaS 운영 툴, 마케팅 실험, 광고 집행", "5,000,000"],
        ["재료비", "현장 테스트 및 시연 자료 제작", "1,000,000"],
        ["지급수수료", "법률, 회계, 사업화 실무 비용", "2,000,000"],
    ]
    for row_idx, values in zip(range(1, 6), budget2_rows):
        fill_row(budget2.rows[row_idx], values)
    fill_row(budget2.rows[6], ["합 계", "합 계", "20,000,000"])

    # Overall roadmap table
    roadmap = tables[8]
    roadmap_rows = [
        ["1", "MVP 고도화", "2026 상반기", "핵심 기능 안정화 및 PoC 준비"],
        ["2", "PoC 운영", "2026 하반기", "외식업 점포 대상 파일럿 검증"],
        ["3", "정식 유료화", "2027 상반기", "Basic 요금제 출시 및 고객 확보"],
        ["4", "Pro 및 제휴 확장", "2027 하반기", "협찬·매칭, B2B 제휴 확대"],
    ]
    for row_idx, values in zip(range(1, 5), roadmap_rows):
        fill_row(roadmap.rows[row_idx], values)

    # Team and partner tables
    team = tables[9]
    team_rows = [
        ["1", "대표자 / AI·CTO", "윤준하 / 서비스 기획, AI 전략, 기술 총괄", "미디어소프트웨어학과 4학년 재학 / AI 구조 설계 및 프로젝트 총괄 역량", "완료"],
        ["2", "공동대표", "노태경 / 마케팅, 기획, 디자인", "대학교 4학년 재학 / 마케팅 기획, 서비스 기획, 디자인 역량", "완료"],
        ["3", "공동대표", "김혜린 / 프론트엔드 개발", "대학교 4학년 재학 / React 기반 UI 구현 및 사용자 경험 설계 역량", "완료"],
        ["4", "공동대표", "오해서 / 백엔드 개발", "대학교 4학년 재학 / Spring Boot 기반 API 및 서버 연동 역량", "완료"],
    ]
    while len(team.rows) < len(team_rows) + 1:
        team.add_row()
    for row_idx, values in zip(range(1, len(team_rows) + 1), team_rows):
        fill_row(team.rows[row_idx], values)

    partners = tables[10]
    partner_rows = [
        ["1", "소상공인시장진흥공단 또는 지역 소상공인지원센터(협력 추진 대상)", "소상공인 지원사업, 정책 연계, 컨설팅 및 현장 네트워크", "외식업 소상공인 대상 파일럿 모집, 정책 연계 자문, 현장 수요 검증", "협약 2~4개월차"],
        ["2", "한국외식업중앙회 지역 지부(협력 추진 대상)", "외식업 종사자 네트워크, 위생교육 및 업계 현장 접점", "파일럿 점포 모집, 업종별 인터뷰, 초기 사용자 확보 및 검증", "협약 3~5개월차"],
        ["3", "예비창업패키지 주관기관 및 창업보육센터(협력 추진 대상)", "BM 고도화, 멘토링, IR 및 사업화 지원", "서비스 사업모델 고도화, 발표평가 대응, 후속 투자 및 판로 연계", "협약 전 기간"],
    ]
    for row_idx, values in zip(range(1, 4), partner_rows):
        fill_row(partners.rows[row_idx], values)

    doc.save(str(OUTPUT))
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(path)
