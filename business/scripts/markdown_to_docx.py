from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DEFAULT_FONT = "Malgun Gothic"


def set_run_font(run, size_pt: float | None = None, bold: bool | None = None) -> None:
    run.font.name = DEFAULT_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph, size_pt: float = 10.5) -> None:
    for run in paragraph.runs:
        set_run_font(run, size_pt=size_pt)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_formatted_runs(paragraph, text: str, size_pt: float = 10.5) -> None:
    text = text.replace("`", "")
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if is_bold else part
        run = paragraph.add_run(clean)
        set_run_font(run, size_pt=size_pt, bold=is_bold)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = DEFAULT_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 20, "111827"),
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13, "1F4E79"),
        ("Heading 3", 11.5, "334155"),
    ]:
        style = styles[style_name]
        style.font.name = DEFAULT_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_summary_box(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EAF2F8")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    add_formatted_runs(
        paragraph,
        "용도: 중간발표 PPT 제작 참고용. 개발 과정의 AI/Agent/Skill 활용, Playwright MCP 기반 검증, 서브 에이전트 역할 분리, 홍보영상 생성 로직의 DSPy 활용을 발표용 문장 중심으로 정리한 문서입니다.",
        size_pt=10,
    )
    doc.add_paragraph()


def add_paragraph(doc: Document, text: str, style: str | None = None, size_pt: float = 10.5):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(4)
    add_formatted_runs(paragraph, text, size_pt=size_pt)
    return paragraph


def convert_markdown_to_docx(input_path: Path, output_path: Path) -> None:
    doc = Document()
    configure_document(doc)

    lines = input_path.read_text(encoding="utf-8").splitlines()
    in_code = False
    first_h1 = True

    for raw_line in lines:
        line = raw_line.rstrip()

        if line.strip().startswith("```"):
            in_code = not in_code
            continue

        if in_code:
            paragraph = add_paragraph(doc, line, size_pt=9.5)
            for run in paragraph.runs:
                run.font.name = "Consolas"
            continue

        if not line.strip():
            continue

        if line.startswith("# "):
            text = line[2:].strip()
            if first_h1:
                paragraph = doc.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_formatted_runs(paragraph, text, size_pt=20)
                doc.add_paragraph()
                add_summary_box(doc)
                first_h1 = False
            else:
                doc.add_heading(text, level=1)
            continue

        heading_match = re.match(r"^(#{2,6})\s+(.+)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            doc.add_heading(heading_match.group(2).strip(), level=level)
            continue

        if line.startswith(">"):
            quote = line.lstrip("> ").strip()
            paragraph = add_paragraph(doc, quote, size_pt=10.5)
            paragraph.paragraph_format.left_indent = Cm(0.45)
            paragraph.paragraph_format.right_indent = Cm(0.2)
            for run in paragraph.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(71, 85, 105)
            continue

        bullet_match = re.match(r"^\s*[-*]\s+(.+)$", line)
        if bullet_match:
            paragraph = add_paragraph(doc, bullet_match.group(1).strip(), style="List Bullet")
            paragraph.paragraph_format.left_indent = Cm(0.55)
            continue

        number_match = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if number_match:
            paragraph = add_paragraph(doc, number_match.group(1).strip(), style="List Number")
            paragraph.paragraph_format.left_indent = Cm(0.55)
            continue

        add_paragraph(doc, line.strip())

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert UTF-8 Markdown to Korean-safe DOCX.")
    parser.add_argument("input", type=Path)
    parser.add_argument("-o", "--output", type=Path)
    args = parser.parse_args()

    output = args.output or args.input.with_suffix(".docx")
    convert_markdown_to_docx(args.input, output)
    print(output.resolve())


if __name__ == "__main__":
    main()
